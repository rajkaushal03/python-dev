from flask import Flask, render_template, request, send_file
import pandas as pd
import os
import re
import docx
import fitz
from spire.doc import *
from spire.doc.common import *


app = Flask(__name__)

def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text
    return full_text

def extract_text_from_doc(file_path):
    doc = Document()
    doc = docx.LoadFromFile (file_path)
    str = doc.GetText()
    return str

def extract_emails(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    return re.findall(email_pattern, text)

def extract_phone_numbers(text):
    phone_pattern = r'\b(?:\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})\b'
    return re.findall(phone_pattern, text)

def process_file(file_path):
    filename, extension = os.path.splitext(file_path)
    if extension == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif extension == '.docx':
        text = extract_text_from_docx(file_path)
    elif extension == '.doc':
        text = extract_text_from_doc(file_path)
    else:
        print(f"Unsupported file format: {extension}")
        return None, None, None
    emails = extract_emails(text)
    phone_numbers = extract_phone_numbers(text)
    return text, emails, phone_numbers

uploads_dir = os.path.join(os.getcwd(), 'uploads')

if not os.path.exists(uploads_dir):
    os.makedirs(uploads_dir)

@app.route('/')
def index():
    return render_template('index.html', data=[], show_download=False)

@app.route('/upload_file', methods=['POST'])
def upload_file():
    files = request.files.getlist('files[]')
    processed_data = []  # Define the processed data list here
    for file in files:
        file_path = os.path.join(uploads_dir, file.filename)
        file.save(file_path)
        try:
            text, emails, phone_numbers = process_file(file_path)
            processed_data.append({'File Name': file.filename, 'Text': text, 'Emails': emails, 'Phone Numbers': phone_numbers})
        except Exception as e:
            print(f"Error processing {file.filename}: {e}")
    if processed_data:
        save_to_excel(processed_data)
        return send_file('cv_data.xlsx', as_attachment=True)
    else:
        return "No data to download."

def save_to_excel(processed_data):
    data_for_excel = []
    for entry in processed_data:
        data_for_excel.append({
            'File Name': entry['File Name'],
            'Text': entry['Text'],
            'Emails': ','.join(entry['Emails']),
            'Phone Numbers': ','.join(entry['Phone Numbers'])
        })
    df = pd.DataFrame(data_for_excel)
    df.to_excel('cv_data.xlsx', index=False)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)
